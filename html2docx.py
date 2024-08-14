from io import BytesIO

from PIL.Image import Image
from bs4 import BeautifulSoup
import requests

from docx import Document


response = requests.get('https://ripe.illinois.edu/blog/difference-between-c3-and-c4-plants')
soup = BeautifulSoup(response.text, 'html.parser')

# print(soup)
article_head = soup.select('article>h4')[0].text
article_p = soup.select('article>div>p')
article_li = soup.select('article>div>ol>li')

print(article_li[0])
print(article_li[1])


doc = Document()

#heading
doc.add_heading(article_head, level=4)
#paragraph
doc.add_paragraph(article_p[0])
doc.add_paragraph(article_p[1])
#list
doc.add_paragraph(f'1. {article_li[0].text}\n2. {article_li[1].text}')

doc.add_paragraph(article_p[2])
doc.add_paragraph(article_p[3])

doc.add_paragraph(article_p[4].text)
doc.add_paragraph(article_p[5])
doc.add_paragraph(article_p[6].text)






doc.save('c3_c4_plants.docx')
